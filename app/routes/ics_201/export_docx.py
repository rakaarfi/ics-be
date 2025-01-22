from fastapi import APIRouter, Depends, HTTPException
from sqlalchemy.ext.asyncio import AsyncSession
from docxtpl import DocxTemplate, InlineImage  # Menggunakan docxtpl untuk templating
from fastapi.responses import FileResponse
from docx.shared import Inches
from datetime import datetime
from sqlmodel import select
import os

from app.config.database import get_session
from app.routes.ics_201.ics_201 import read_ics_201_by_id
from app.routes.ics_201.resource_summary import read_resource_summary_by_ics_201_id
from app.routes.ics_201.actions_strategies_tactics import read_actions_strategies_tactics_by_ics_id
from app.routes.ics_201.chart import read_ics_chart_by_ics_id
from app.routes.incident_data import read_incident_data_by_id
from app.models.imt_members.main_section import *
from app.models.imt_members.finance_section import *
from app.models.imt_members.logistic_section import *
from app.models.imt_members.planning_section import *
from app.models.ics_201 import IcsChartBase, IcsChart

router = APIRouter()

# async def generate_docx(
#     template_path: str, 
#     output_path: str, 
#     context: dict,
#     image_directory: str
# ):
#     doc = Document(template_path)
    
#     # Ganti teks placeholder dengan nilai dari context, kecuali {{ map_sketch }}
#     for paragraph in doc.paragraphs:
#         for key, value in context.items():
#             if key != "map_sketch":  # Jangan ganti {{ map_sketch }} dengan teks
#                 placeholder = f"{{{{ {key} }}}}"
#                 if placeholder in paragraph.text:
#                     paragraph.text = paragraph.text.replace(placeholder, str(value))
                
#     for table in doc.tables:
#         for row in table.rows:
#             for cell in row.cells:
#                 for key, value in context.items():
#                     if key != "map_sketch":  # Jangan ganti {{ map_sketch }} dengan teks
#                         placeholder = f"{{{{ {key} }}}}"
#                         if placeholder in cell.text:
#                             cell.text = cell.text.replace(placeholder, str(value))
    
#     # Tambahkan gambar ke dokumen jika ada kolom map_sketch
#     if "map_sketch" in context and context["map_sketch"]:
#         image_filename = context["map_sketch"]
#         image_path = os.path.join(image_directory, image_filename)
        
#         if os.path.exists(image_path):
#             # Cari placeholder {{ map_sketch }} di paragraf
#             for paragraph in doc.paragraphs:
#                 placeholder = "{{ map_sketch }}"
#                 if placeholder in paragraph.text:
#                     print(f"Placeholder ditemukan dalam paragraf: {paragraph.text}")
#                     paragraph.text = paragraph.text.replace(placeholder, "")
#                     run = paragraph.add_run()
#                     run.add_picture(image_path, width=Inches(4.0))  # Sesuaikan ukuran gambar
            
#             # Cari placeholder {{ map_sketch }} di tabel
#             for table in doc.tables:
#                 for row in table.rows:
#                     for cell in row.cells:
#                         placeholder = "{{ map_sketch }}"
#                         if placeholder in cell.text:
#                             print(f"Placeholder ditemukan dalam cell: {cell.text}")
#                             cell.text = cell.text.replace(placeholder, "")
#                             # Akses paragraf di dalam sel dan tambahkan gambar
#                             paragraph = cell.paragraphs[0]  # Ambil paragraf pertama di sel
#                             run = paragraph.add_run()
#                             run.add_picture(image_path, width=Inches(4.0))  # Sesuaikan ukuran gambar
#         else:
#             print(f"File gambar {image_path} tidak ditemukan.")
            
#     # Bagian kode untuk actions_strategies_tactics
#     if "actions_strategies_tactics" in context and context["actions_strategies_tactics"]:
#         # print("Contoh data actions_strategies_tactics:", context["actions_strategies_tactics"][0])
        
#         target_table_index = 1  # Tabel ke-2
#         if len(doc.tables) > target_table_index:
#             table = doc.tables[target_table_index]
            
#             # Cari baris yang memiliki placeholder
#             start_row_index = None
#             for row_idx, row in enumerate(table.rows):
#                 for cell in row.cells:
#                     # Periksa setiap paragraf dalam sel
#                     for paragraph in cell.paragraphs:
#                         if "{{" in paragraph.text and "actions_strategies_tactics" in paragraph.text and "}}" in paragraph.text:
#                             start_row_index = row_idx
#                             break
#                 if start_row_index is not None:
#                     break

#             if start_row_index is not None:
#                 # print(f"Placeholder ditemukan di baris {start_row_index}")
#                 current_row = start_row_index
                
#                 # Isi data ke dalam tabel
#                 for rs in context["actions_strategies_tactics"]:
#                     # Jika masih ada baris tersedia
#                     if current_row < len(table.rows):
#                         row = table.rows[current_row]
                        
#                         # Hapus placeholder dari sel pertama
#                         if current_row == start_row_index:
#                             for cell in row.cells:
#                                 if "{{actions_strategies_tactics}}" in cell.text:
#                                     cell.text = ""
                        
#                         # Isi data ke setiap kolom
#                         try:
#                             # Kolom 1: Jam
#                             if len(row.cells) > 0:
#                                 row.cells[0].text = rs['time_initiated'].strftime('%H:%M') if rs.get('time_initiated') else ''
                            
#                             # Kolom 2: Tindakan
#                             if len(row.cells) > 1:
#                                 row.cells[1].text = str(rs.get("actions", ""))
                            
#                             current_row += 1
                            
#                         except Exception as e:
#                             print(f"Error saat mengisi baris {current_row}: {str(e)}")
#                     else:
#                         # Jika tidak ada baris tersedia, tambahkan baris baru
#                         new_row = table.add_row()
#                         try:
#                             new_row.cells[0].text = rs['time_initiated'].strftime('%H:%M') if rs.get('time_initiated') else ''
#                             new_row.cells[1].text = str(rs.get("actions", ""))
#                         except Exception as e:
#                             print(f"Error saat menambah baris baru: {str(e)}")
#             else:
#                 print("Placeholder {{actions_strategies_tactics}} tidak ditemukan dalam tabel")
#         else:
#             print(f"Tabel ke-{target_table_index + 1} tidak ditemukan.")
            
#     # Bagian kode untuk resource_summaries
#     if "resource_summaries" in context and context["resource_summaries"]:
#         # print("Contoh data resource_summaries:", context["resource_summaries"][0])
        
#         target_table_index = 3  # Tabel ke-4
#         if len(doc.tables) > target_table_index:
#             table = doc.tables[target_table_index]
            
#             # Cari baris yang memiliki placeholder
#             start_row_index = None
#             for row_idx, row in enumerate(table.rows):
#                 for cell in row.cells:
#                     # Periksa setiap paragraf dalam sel
#                     for paragraph in cell.paragraphs:
#                         if "{{" in paragraph.text and "resource_summaries" in paragraph.text and "}}" in paragraph.text:
#                             start_row_index = row_idx
#                             break
#                 if start_row_index is not None:
#                     break

#             if start_row_index is not None:
#                 print(f"Placeholder ditemukan di baris {start_row_index}")
#                 current_row = start_row_index
                
#                 # Isi data ke dalam tabel
#                 for rs in context["resource_summaries"]:
#                     # Jika masih ada baris tersedia
#                     if current_row < len(table.rows):
#                         row = table.rows[current_row]
                        
#                         # Hapus placeholder dari sel pertama
#                         if current_row == start_row_index:
#                             for cell in row.cells:
#                                 if "{{resource_summaries}}" in cell.text:
#                                     cell.text = ""
                        
#                         # Isi data ke setiap kolom
#                         try:
#                             # Kolom 1: Sumber Daya
#                             if len(row.cells) > 0:
#                                 row.cells[0].text = str(rs.get("resource", ""))
                            
#                             # Kolom 2: Pengidentifikasi Sumber Daya
#                             if len(row.cells) > 1:
#                                 row.cells[1].text = str(rs.get("resource_identified", ""))
                            
#                             # Kolom 3: Tanggal/Jam Dipesan
#                             if len(row.cells) > 2:
#                                 date_str = rs['date_ordered'].strftime('%Y-%m-%d') if rs.get('date_ordered') else ''
#                                 time_str = rs['time_ordered'].strftime('%H:%M') if rs.get('time_ordered') else ''
#                                 row.cells[2].text = f"{date_str} {time_str}"
                            
#                             # Kolom 5: ETA (indeks 4 karena kolom 4 kosong)
#                             if len(row.cells) > 4:
#                                 row.cells[4].text = str(rs.get("eta", ""))
                            
#                             # Kolom 6: Status Tiba
#                             if len(row.cells) > 5:
#                                 row.cells[5].text = "Tiba" if rs.get("is_arrived", False) else "Belum Tiba"
                            
#                             # Kolom 7: Catatan
#                             if len(row.cells) > 6:
#                                 row.cells[6].text = str(rs.get("notes", ""))
                            
#                             current_row += 1
                            
#                         except Exception as e:
#                             print(f"Error saat mengisi baris {current_row}: {str(e)}")
#                     else:
#                         # Jika tidak ada baris tersedia, tambahkan baris baru
#                         new_row = table.add_row()
#                         try:
#                             new_row.cells[0].text = str(rs.get("resource", ""))
#                             new_row.cells[1].text = str(rs.get("resource_identified", ""))
#                             date_str = rs['date_ordered'].strftime('%Y-%m-%d') if rs.get('date_ordered') else ''
#                             time_str = rs['time_ordered'].strftime('%H:%M') if rs.get('time_ordered') else ''
#                             new_row.cells[2].text = f"{date_str} {time_str}"
#                             new_row.cells[4].text = str(rs.get("eta", ""))
#                             new_row.cells[5].text = "Tiba" if rs.get("is_arrived", False) else "Belum Tiba"
#                             new_row.cells[6].text = str(rs.get("notes", ""))
#                         except Exception as e:
#                             print(f"Error saat menambah baris baru: {str(e)}")
#             else:
#                 print("Placeholder {{resource_summaries}} tidak ditemukan dalam tabel")
#         else:
#             print(f"Tabel ke-{target_table_index + 1} tidak ditemukan.")
            
#     doc.save(output_path)

# async def generate_docx(
#     template_path: str, 
#     output_path: str, 
#     context: dict,
#     image_directory: str
# ):
#     """
#     Fungsi untuk menghasilkan dokumen Word (.docx) dari template.
#     Menggunakan docxtpl untuk mengganti placeholder di text box dan elemen lainnya.
#     """
#     # Muat template
#     doc = DocxTemplate(template_path)

#     # Kalau kolom map_sketch ada isinya, jadikan InlineImage
#     if map_sketch_filename:
#         image_path = os.path.join(image_directory, map_sketch_filename)
#         if os.path.exists(image_path):
#             context["map_sketch"] = InlineImage(doc, image_path, width=Inches(4.0))
#         else:
#             # Bila file gambar tidak ditemukan, kosongkan saja atau beri tanda
#             context["map_sketch"] = "Gambar tidak ditemukan"

#     # Render docxtpl
#     doc.render(context)

#     # Simpan hasil
#     doc.save(output_path)



def format_date_to_sentence(date_str):
    # Parse tanggal dari string ke objek datetime
    date_obj = datetime.strptime(date_str, "%Y-%m-%d")
    
    # Ambil hari, bulan, dan tahun
    day = date_obj.day
    month = date_obj.strftime("%B")  # Nama bulan (e.g., January)
    year = date_obj.year
    
    # Tambahkan suffix untuk hari
    if 4 <= day <= 20 or 24 <= day <= 30:
        suffix = "th"
    else:
        suffix = ["st", "nd", "rd"][day % 10 - 1]
    
    # Gabungkan menjadi kalimat
    return f"{day}{suffix} {month} {year}"

@router.post("/export_docx/{id}")
async def export_docx(id: int, session: AsyncSession = Depends(get_session)):
    template_path = "app/docx_template/ics_201.docx"
    output_path = f"app/docx_template/ics_201_{id}.docx"
    
    # Ambil data Ics201
    ics_201 = await read_ics_201_by_id(id, session)
    context = ics_201.dict()
    context['date_initiated'] = format_date_to_sentence(context['date_initiated'].strftime('%Y-%m-%d'))
    
    incident_id = context['incident_id']
    # Ambil incident_data yang terkait
    incident_data = await read_incident_data_by_id(id=incident_id, session=session)
    context["incident_data"] = incident_data
    context["incident_name"] = incident_data.name
    context["incident_no"] = incident_data.no
    
    # Ambil data resource_summary yang terkait
    resource_summaries = await read_resource_summary_by_ics_201_id(id, session)
    resource_summaries_list = [rs.dict() for rs in resource_summaries]
    
    # Manipulasi is_arrived
    for rs in resource_summaries_list:
        rs['is_arrived'] = '✓' if rs['is_arrived'] else '✗'
        rs['date_ordered'] = format_date_to_sentence(rs['date_ordered'].strftime('%Y-%m-%d'))
        rs['time_ordered'] = rs['time_ordered'].strftime('%H:%M')
    
    context["resource_summaries"] = resource_summaries_list
    
    # Ambil data actions_strategies_tactics yang terkait
    actions_strategies_tactics = await read_actions_strategies_tactics_by_ics_id(id, session)
    actions_strategies_tactics_list = [ast.dict() for ast in actions_strategies_tactics]
    
    # Manipulasi time_initiated
    for ast in actions_strategies_tactics_list:
        ast['time_initiated'] = ast['time_initiated'].strftime('%H:%M')
    
    context["actions_strategies_tactics"] = actions_strategies_tactics_list
    
    # Ambil data chart yang terkait
    charts = await read_ics_chart_by_ics_id(id, session)
    if charts:
        chart = charts[0]  # Ambil chart pertama dari list
        chart_data = await read_docx_chart(chart.id, session)
        context.update(chart_data.dict())  # Tambahkan data nama ke context
    else:
        raise HTTPException(status_code=404, detail="Chart tidak ditemukan")
    
    # Direktori tempat gambar disimpan
    image_directory = "uploaded_files"
    map_sketch_filename = context.get("map_sketch")
    
    # Muat template
    doc = DocxTemplate(template_path)

    # Kalau kolom map_sketch ada isinya, jadikan InlineImage
    if map_sketch_filename:
        image_path = os.path.join(image_directory, map_sketch_filename)
        if os.path.exists(image_path):
            context["map_sketch"] = InlineImage(doc, image_path, width=Inches(4.0))
        else:
            # Bila file gambar tidak ditemukan, kosongkan saja atau beri tanda
            context["map_sketch"] = "Gambar tidak ditemukan"

    # Render docxtpl
    doc.render(context)

    # Simpan hasil
    doc.save(output_path)
    
    return FileResponse(output_path, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', filename=f"ics_201_{id}.docx")
    
async def read_docx_chart(
    id: int,
    session: AsyncSession = Depends(get_session)
) -> IcsChartBase:
    # Query construction remains the same
    statement = (
        select(
            IncidentCommander.name.label('incident_commander_name'),
            DeputyIncidentCommander.name.label('deputy_incident_commander_name'),
            SafetyOfficer.name.label('safety_officer_name'),
            PublicInformationOfficer.name.label('public_information_officer_name'),
            LiaisonOfficer.name.label('liaison_officer_name'),
            LegalOfficer.name.label('legal_officer_name'),
            HumanCapitalOfficer.name.label('human_capital_officer_name'),
            OperationSectionChief.name.label('operation_section_chief_name'),
            PlanningSectionChief.name.label('planning_section_chief_name'),
            SituationUnitLeader.name.label('situation_unit_leader_name'),
            ResourcesUnitLeader.name.label('resources_unit_leader_name'),
            DocumentationUnitLeader.name.label('documentation_unit_leader_name'),
            DemobilizationUnitLeader.name.label('demobilization_unit_leader_name'),
            EnvironmentalUnitLeader.name.label('environmental_unit_leader_name'),
            TechnicalSpecialist.name.label('technical_specialist_name'),
            LogisticSectionChief.name.label('logistic_section_chief_name'),
            CommunicationUnitLeader.name.label('communication_unit_leader_name'),
            MedicalUnitLeader.name.label('medical_unit_leader_name'),
            FoodUnitLeader.name.label('food_unit_leader_name'),
            FacilityUnitLeader.name.label('facility_unit_leader_name'),
            SupplyUnitLeader.name.label('supply_unit_leader_name'),
            TransportationUnitLeader.name.label('transportation_unit_leader_name'),
            FinanceSectionChief.name.label('finance_section_chief_name'),
            ProcurementUnitLeader.name.label('procurement_unit_leader_name'),
            CompensationClaimUnitLeader.name.label('compensation_claim_unit_leader_name'),
            CostUnitLeader.name.label('cost_unit_leader_name'),
            TimeUnitLeader.name.label('time_unit_leader_name'),
        )
        .join(IcsChart.incident_commander)
        .join(IcsChart.deputy_incident_commander)
        .join(IcsChart.safety_officer)
        .join(IcsChart.public_information_officer)
        .join(IcsChart.liaison_officer)
        .join(IcsChart.legal_officer)
        .join(IcsChart.human_capital_officer)
        .join(IcsChart.operation_section_chief)
        .join(IcsChart.planning_section_chief)
        .join(IcsChart.situation_unit_leader)
        .join(IcsChart.resources_unit_leader)
        .join(IcsChart.documentation_unit_leader)
        .join(IcsChart.demobilization_unit_leader)
        .join(IcsChart.environmental_unit_leader)
        .join(IcsChart.technical_specialist)
        .join(IcsChart.logistic_section_chief)
        .join(IcsChart.communication_unit_leader)
        .join(IcsChart.medical_unit_leader)
        .join(IcsChart.food_unit_leader)
        .join(IcsChart.facility_unit_leader)
        .join(IcsChart.supply_unit_leader)
        .join(IcsChart.transportation_unit_leader)
        .join(IcsChart.finance_section_chief)
        .join(IcsChart.procurement_unit_leader)
        .join(IcsChart.compensation_claim_unit_leader)
        .join(IcsChart.cost_unit_leader)
        .join(IcsChart.time_unit_leader)
        .where(IcsChart.id == id)
    )
    
    result = await session.execute(statement)
    data = result.first()
    
    if not data:
        raise HTTPException(status_code=404, detail="Item not found")
    
    return IcsChartBase(**data._asdict())